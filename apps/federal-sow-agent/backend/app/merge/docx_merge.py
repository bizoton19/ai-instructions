"""Merge generated SOW fields into Word templates."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from app.schemas import SOWSectionsModel


def _has_jinja_placeholders(path: Path) -> bool:
    try:
        doc = Document(str(path))
        pattern = re.compile(r"\{\{\s*[\w.]+\s*\}\}")
        for p in doc.paragraphs:
            if pattern.search(p.text):
                return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if pattern.search(p.text):
                            return True
        return False
    except Exception:
        return True


def merge_docx(
    template_path: Path,
    flat_context: dict[str, str],
    output_path: Path,
) -> tuple[bool, str]:
    """
    Returns (used_docxtpl, message).
    If docxtpl fails or no placeholders detected, appends sections as new document content.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    used_tpl = False
    msg_parts: list[str] = []

    if _has_jinja_placeholders(template_path):
        try:
            tpl = DocxTemplate(str(template_path))
            tpl.render(flat_context)
            tpl.save(str(output_path))
            used_tpl = True
            msg_parts.append("Rendered template with docxtpl.")
            return used_tpl, " ".join(msg_parts)
        except Exception as e:
            msg_parts.append(f"docxtpl render failed ({e}); using fallback.")

    doc = Document(str(template_path))
    doc.add_paragraph()
    doc.add_heading("Generated Statement of Work Content", level=1)
    order = [
        ("Purpose", flat_context.get("purpose", "")),
        ("Background", flat_context.get("background", "")),
        ("Scope", flat_context.get("scope", "")),
        ("Deliverables", flat_context.get("deliverables", "")),
        ("Period of Performance", flat_context.get("period_of_performance", "")),
        ("Roles and Responsibilities", flat_context.get("roles_and_responsibilities", "")),
        ("Acceptance Criteria", flat_context.get("acceptance_criteria", "")),
        ("Assumptions and Constraints", flat_context.get("assumptions_and_constraints", "")),
    ]
    body = flat_context.get("full_markdown") or ""
    for title, block in order:
        if block and str(block).strip():
            doc.add_heading(title, level=2)
            for line in str(block).splitlines():
                doc.add_paragraph(line)
    if body.strip():
        doc.add_heading("Full narrative (Markdown pasted as text)", level=2)
        for line in body.splitlines():
            doc.add_paragraph(line)

    doc.save(str(output_path))
    msg_parts.append("Fallback: appended generated sections to a copy of the template.")
    return used_tpl, " ".join(msg_parts)


def merge_or_standalone_docx(
    template_path: Path,
    template_filename: str,
    flat_context: dict[str, str],
    output_path: Path,
) -> str:
    """
    If the workspace template is .docx, merge/render into that file; otherwise build a new .docx
    guided by PDF/XLSX template (same rules as HTTP merge/export).
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = Path(template_filename).suffix.lower()
    if suffix == ".docx":
        _, note = merge_docx(template_path, flat_context, output_path)
        return note
    preamble = (
        "Reference template file type: PDF or Excel. The system extracted headings, text, or table "
        "previews for the drafting specialists. This Word file contains the generated Statement of Work "
        "content in structured sections—it does not recreate the original file’s exact layout, forms, or print styling."
    )
    return standalone_docx_from_flat(flat_context, output_path, preamble=preamble)


def standalone_docx_from_flat(
    flat_context: dict[str, str],
    output_path: Path,
    *,
    preamble: str | None = None,
) -> str:
    """
    Word document from structured sections only (no template file).
    PDF/Excel uploads can inform drafting via extracted hints but cannot merge with python-docx.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if preamble and preamble.strip():
        doc.add_paragraph(preamble.strip())
    doc.add_paragraph()
    doc.add_heading("Generated Statement of Work Content", level=1)
    order = [
        ("Purpose", flat_context.get("purpose", "")),
        ("Background", flat_context.get("background", "")),
        ("Scope", flat_context.get("scope", "")),
        ("Deliverables", flat_context.get("deliverables", "")),
        ("Period of Performance", flat_context.get("period_of_performance", "")),
        ("Roles and Responsibilities", flat_context.get("roles_and_responsibilities", "")),
        ("Acceptance Criteria", flat_context.get("acceptance_criteria", "")),
        ("Assumptions and Constraints", flat_context.get("assumptions_and_constraints", "")),
    ]
    body = flat_context.get("full_markdown") or ""
    for title, block in order:
        if block and str(block).strip():
            doc.add_heading(title, level=2)
            for line in str(block).splitlines():
                doc.add_paragraph(line)
    if body.strip():
        doc.add_heading("Full narrative (Markdown pasted as text)", level=2)
        for line in body.splitlines():
            doc.add_paragraph(line)

    doc.save(str(output_path))
    return "Created new Word document from generated sections (no .docx template binary used as merge target)."


def merge_docx_bytes(template_path: Path, flat_context: dict[str, str]) -> tuple[bytes, str]:
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = Path(tmp.name)
    try:
        _, msg = merge_docx(template_path, flat_context, out_path)
        return out_path.read_bytes(), msg
    finally:
        out_path.unlink(missing_ok=True)


def sow_sections_to_markdown(s: SOWSectionsModel) -> str:
    """Produce a Markdown string from structured sections (prefers ``full_markdown``)."""

    fm = (s.full_markdown or "").strip()
    if fm:
        return fm

    parts: list[str] = []
    order = [
        ("Purpose", "purpose"),
        ("Background", "background"),
        ("Scope", "scope"),
        ("Deliverables", "deliverables"),
        ("Period of Performance", "period_of_performance"),
        ("Roles and Responsibilities", "roles_and_responsibilities"),
        ("Acceptance Criteria", "acceptance_criteria"),
        ("Assumptions and Constraints", "assumptions_and_constraints"),
    ]
    data = s.model_dump()
    for title, key in order:
        block = (data.get(key) or "").strip()
        if block:
            parts.append(f"## {title}\n\n{block}\n")
    out = "\n".join(parts).strip()
    if out:
        return out
    return "[No generated narrative text was available yet.]"


def sow_model_to_flat(s: SOWSectionsModel) -> dict[str, str]:
    return {
        "purpose": s.purpose or "",
        "background": s.background or "",
        "scope": s.scope or "",
        "deliverables": s.deliverables or "",
        "period_of_performance": s.period_of_performance or "",
        "roles_and_responsibilities": s.roles_and_responsibilities or "",
        "acceptance_criteria": s.acceptance_criteria or "",
        "assumptions_and_constraints": s.assumptions_and_constraints or "",
        "full_markdown": s.full_markdown or "",
    }
