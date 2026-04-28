"""Extract text and heading outline from .docx."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


def extract_docx_text_and_outline(path: Path) -> tuple[str, list[str]]:
    doc = Document(str(path))
    lines: list[str] = []
    outline: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        lines.append(text)
        style = para.style.name if para.style else ""
        if style.startswith("Heading") or style.startswith("heading"):
            outline.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    full_text = "\n".join(lines)
    return full_text, outline


def outline_to_json(outline: list[str]) -> str:
    return json.dumps({"headings": outline})
